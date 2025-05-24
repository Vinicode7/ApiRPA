import sqlite3
from datetime import datetime
from docx import Document
from docx.shared import Inches

def buscar_dados_paises():
    conexao = sqlite3.connect("dados/paises.db")
    cursor = conexao.cursor()
    cursor.execute("SELECT * FROM paises")
    resultados = cursor.fetchall()
    conexao.close()
    return resultados

def buscar_dados_livros():
    conexao = sqlite3.connect("dados/livraria.db")
    cursor = conexao.cursor()
    cursor.execute("SELECT * FROM livros")
    resultados = cursor.fetchall()
    conexao.close()
    return resultados

def gerar_relatorio(nome_aluno):
    doc = Document()
    doc.add_heading('Relatório Final - Países e Livros', 0)

    data_atual = datetime.now().strftime("%d/%m/%Y %H:%M")
    doc.add_paragraph(f"Aluno: {nome_aluno}")
    doc.add_paragraph(f"Data de geração: {data_atual}")

    # Tabela Países
    doc.add_heading('Dados dos Países', level=1)
    paises = buscar_dados_paises()
    if paises:
        tabela_paises = doc.add_table(rows=1, cols=13)
        hdr_cells = tabela_paises.rows[0].cells
        headers = ["Nome comum", "Nome oficial", "Capital", "Continente", "Região", "Sub-região",
                   "População", "Área", "Moeda (nome)", "Moeda (símbolo)", "Idioma principal",
                   "Fuso horário", "URL bandeira"]
        for i, h in enumerate(headers):
            hdr_cells[i].text = h

        for p in paises:
            row_cells = tabela_paises.add_row().cells
            for i, valor in enumerate(p):
                row_cells[i].text = str(valor)
    else:
        doc.add_paragraph("Nenhum dado de país encontrado.")

    doc.add_paragraph()  

    # aq fiz a tabela Livros
    doc.add_heading('Dados dos Livros', level=1)
    livros = buscar_dados_livros()
    if livros:
        tabela_livros = doc.add_table(rows=1, cols=4)
        hdr_cells = tabela_livros.rows[0].cells
        headers = ["Título", "Preço", "Avaliação", "Disponibilidade"]
        for i, h in enumerate(headers):
            hdr_cells[i].text = h

        for l in livros:
            row_cells = tabela_livros.add_row().cells
            for i, valor in enumerate(l):
                row_cells[i].text = str(valor)
    else:
        doc.add_paragraph("Nenhum dado de livro encontrado.")

    # code para salvar arquivo
    nome_arquivo = f"relatorio_final_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    doc.save(nome_arquivo)
    print(f"✅ Relatório gerado com sucesso: {nome_arquivo}")

if __name__ == "__main__":
    nome = input("Digite seu nome para o relatório: ")
    gerar_relatorio(nome)
